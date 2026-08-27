
import hashlib
from typing import List, Tuple, Dict

import chromadb
from pathlib import Path
import json
from docx import Document as DocxDocument

from pypdf import PdfReader

from app.rag_pipeline.dense_search import VectorStore
from app.rag_pipeline.fiuse import reciprocal_rank_fusion
from app.rag_pipeline.reranker import Reranker
from app.rag_pipeline.schema import Document, TextChunk
from app.rag_pipeline.sparce_search import BM25Index
from app.rag_pipeline.preprocessing import chunk_document, clean_text, DEFAULT_CHUNK_SIZE, DEFAULT_OVERLAP, DEFAULT_TOP_K, DEFAULT_TOP_N


class RAG:
    def __init__(self, persist_dir: str = "./rag_store", chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_OVERLAP, top_k_retrieve: int = DEFAULT_TOP_K, top_n_rerank: int = DEFAULT_TOP_N):
    
        self.persist_dir = Path(persist_dir)
        self.persist_dir.mkdir(exist_ok=True)

        self.chunk_size = chunk_size
        self.overlap = overlap
        self.top_k_retrieve = top_k_retrieve
        self.top_n_rerank = top_n_rerank

     
        self.vector_store = VectorStore(persist_dir=str(self.persist_dir / "chroma"))
        self.bm25_index = BM25Index()
        self.reranker = Reranker()

        self._all_chunks: List[TextChunk]=[]
        self._ingested_sources: Dict[str, dict] = {}

    def ingest(self, file_path:str):
        print (f"Starting ingestion for {file_path} ...")
        source_name=Path(file_path).name

        if source_name in self._ingested_sources:
            print (f"Source {source_name} already ingested. Skipping.")
            return {**self._ingested_sources[source_name], "status": "already_ingested"}

        chunks_path = self.persist_dir / f"{source_name}_chunks.json"
        stats:dict ={"source":source_name}
        print (f"Checking for existing chunks at {chunks_path} ...")


        if chunks_path.exists():
            print(f"Loading existing chunks for {source_name} from {chunks_path}")
            source_chunks = self._load_chunks(chunks_path)
            stats["cached"]=True

        else:
            print(f"processing {source_name} ...")

            raw_text = self.load_text_from_file(file_path)

            cleaned_text = clean_text(raw_text)
            print(f"[ingest] Cleaned text length: {len(cleaned_text)} characters", flush=True)
            source_chunks = chunk_document(
                Document(text=cleaned_text, source=source_name, metadata={},document_id=source_name),
                chunk_size=self.chunk_size,
                overlap=self.overlap
            )

            self._save_chunks(source_chunks, chunks_path)
            self.vector_store.index_chunks(source_chunks, source_id=source_name)

        stats["chunks"]=len(source_chunks)
        self._all_chunks.extend(source_chunks)
        self.bm25_index.build(self._all_chunks)
        self._ingested_sources[source_name]=stats

        stats["total_chunks"]=len(self._all_chunks)
        stats["total_sources"]=len(self._ingested_sources)
        print(f"Ingestion complete for {source_name}. Total chunks: {len(self._all_chunks)}. Total sources: {len(self._ingested_sources)}")

        self._save_state()
        return stats

    def ingest_multiple(self, file_paths: List[str]):
        for file_path in file_paths:
            self.ingest(file_path)

    def _save_chunks(self, chunks: List[TextChunk], path: Path):
        with open(path, "w", encoding="utf-8") as f:
            json.dump([chunk.__dict__ for chunk in chunks], f, ensure_ascii=False, indent=4)   
    
    def _load_chunks(self, path: Path) -> List[TextChunk]:
        with open(path, "r", encoding="utf-8") as f:
            chunks_data = json.load(f)  

        return [TextChunk(**data) for data in chunks_data]
    def load_text_from_file(self, file_path: str) -> str:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        extension = path.suffix.lower()


        if extension == ".txt":

            return path.read_text(
                encoding="utf-8",
                errors="replace"
            )
        elif extension == ".pdf":

            reader = PdfReader(file_path)

            pages = []

            for page in reader.pages:

                page_text = page.extract_text() or ""

                if page_text.strip():
                    pages.append(page_text)

            return "\n\n".join(pages)

        elif extension == ".docx":

            document = DocxDocument(file_path)

            paragraphs = []

            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if text:
                    paragraphs.append(text)

            return "\n".join(paragraphs)

        
        else:

            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: .txt, .pdf, .docx"
            )
    
    def retrieve(self, query: str) -> List[Tuple[str, dict, float]]:
        print(f"[retrieve] Starting query: {query!r}", flush=True)
        dense_results = self.vector_store.dense_search(query, top_k=self.top_k_retrieve)
        print(f"[retrieve] Dense results: {len(dense_results)}", flush=True)
        bm25_results = self.bm25_index.search(query, top_k=self.top_k_retrieve)
        print(f"[retrieve] BM25 results: {len(bm25_results)}", flush=True)

        fused = reciprocal_rank_fusion(dense_results, bm25_results)
        print(f"[retrieve] Fused results: {len(fused)}", flush=True)
        if not fused :
            print("[retrieve] No fused results; stopping", flush=True)
            return []
        reranked_results = self.reranker.rerank(query, fused, top_n=self.top_n_rerank)
        print(f"[retrieve] Reranked results: {len(reranked_results)}", flush=True)
        return reranked_results

    
    def _save_state(self):
        print(f"Saving RAG state to {self.persist_dir / 'state.json'} ...")
        state={"ingested_sources": list(self._ingested_sources.keys())}
        ( self.persist_dir / "state.json").write_text(json.dumps(state, indent=4))

    def format_context(self, hits):
        context_segments = []

        for text, metadata, score in hits:
            source = metadata.get("source", "")
            chunk_id = metadata.get("chunk_id", "")

            segment = (
                f"[Source: {source} | Chunk: {chunk_id}]\n"
                f"{text}"
            )

            context_segments.append(segment)

        return "\n\n".join(context_segments)



    def load_state(self) -> bool:

        state_path = self.persist_dir / "state.json"
        print(f"[state] Loading state from {state_path}", flush=True)
        if not state_path.exists():
            print("[state] No state file found", flush=True)
            return False
        try:

            state = json. loads(state_path. read_text())
            sources = state.get("ingested_sources", [])

            for name in sources:
                chunks_path = self.persist_dir / f"{name}_chunks.json"
                print(f"[state] Checking cached chunks for {name!r}: {chunks_path}", flush=True)
                if not chunks_path.exists():
                    print(f"[state] Cached chunks missing for {name!r}", flush=True)
                    continue
                source_chunks = self ._load_chunks(chunks_path)
                self ._all_chunks.extend(source_chunks)
                total = len(self ._all_chunks)
                self ._ingested_sources [name] = {
                "source":name,
                "chunks":len(source_chunks),
                "cached":True,
                "total_chunks": total,
                "total_sources": 0 # updated below
                }

            if self._all_chunks:
                self.bm25_index.build(self ._all_chunks)
                # Fix total_sources counts
                n = len(self ._ingested_sources)
                for v in self ._ingested_sources.values():
                    v["total_sources"] = n


            print(f" Restored {len(self ._ingested_sources)} source(s) from cache.")
            return bool(self ._ingested_sources)

        except Exception as e:
            print(f"A Could not restore RAG state: {e}")
        return False

        

    def _save_source_chunks(self, chunks: List[TextChunk], path: Path):
        data = [
            {
                "chunk_id": chunk.chunk_id,
                "text": chunk.text,
                "source": chunk.source,
                "metadata": chunk.metadata,
                "char_start": chunk.char_start,
                "char_end": chunk.char_end,
                "token_count": chunk.token_count,
            }
            for chunk in chunks
        ]

        path.write_text(json.dumps(data, ensure_ascii=False, indent=2))
